import streamlit as st
import pandas as pd
import re
import unicodedata
from io import BytesIO
from docx import Document

st.title("Normalizador Word y Excel")

# -------------------------
# NORMALIZAR TEXTO
# -------------------------
def normalizar(texto):
    texto = str(texto).strip().lower()
    texto = unicodedata.normalize('NFKD', texto)
    texto = texto.encode('ascii', 'ignore').decode('utf-8')
    texto = re.sub(r"[ .\-\/]+", "_", texto)
    texto = re.sub(r"[^a-z0-9_]", "", texto)
    texto = re.sub(r"_+", "_", texto)
    return texto.strip("_")

# -------------------------
# PROCESAR PÁRRAFO SIN ROMPER FORMATO
# -------------------------
def procesar_parrafo(paragraph):

    # unir texto completo del párrafo
    full_text = "".join(run.text for run in paragraph.runs)

    # si no hay variables → salir
    if "{{" not in full_text:
        return

    # buscar variables completas
    variables = re.findall(r"{{(.*?)}}", full_text)

    for var in variables:
        nueva = normalizar(var)
        full_text = full_text.replace(
            "{{" + var + "}}",
            "{{" + nueva + "}}"
        )

    # ✅ reescribir SOLO si hubo cambio
    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""

# -------------------------
# NORMALIZAR WORD COMPLETO
# -------------------------
def normalizar_word(doc):

    # párrafos normales
    for para in doc.paragraphs:
        procesar_parrafo(para)

    # tablas (CLAVE en tu plantilla)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    procesar_parrafo(para)

# -------------------------
# SUBIR ARCHIVOS
# -------------------------
word_file = st.file_uploader("📄 Subir Plantilla Word", type=["docx"])
excel_file = st.file_uploader("📊 Subir data Excel", type=["xlsx"])

# -------------------------
# PROCESAR
# -------------------------
if word_file and excel_file:

    # ✅ PROCESAR WORD
    doc = Document(word_file)
    normalizar_word(doc)

    word_buffer = BytesIO()
    doc.save(word_buffer)

    # ✅ LIMPIAR EXCEL
    df = pd.read_excel(excel_file, dtype=str)

    # normalizar columnas
    df.columns = [normalizar(col) for col in df.columns]

    # eliminar filas vacías
    df = df.dropna(how="all")

    # 🔥 regla clave -> eliminar si nro vacío o 0
    if "nro" in df.columns:
        df["nro"] = df["nro"].fillna("").astype(str).str.strip()
        df = df[df["nro"] != ""]
        df = df[df["nro"] != "0"]

    # convertir todo a texto
    df = df.fillna("").astype(str)
    df = df.apply(lambda col: col.str.strip())

    # preview
    st.subheader("✅ Datos limpios")
    st.dataframe(df.head())

    # exportar excel
    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False)

    # -------------------------
    # DESCARGAS
    # -------------------------
    st.download_button(
        "📄 Descargar Word Normalizado",
        data=word_buffer.getvalue(),
        file_name="word_normalizado.docx"
    )

    st.download_button(
        "📊 Descargar Excel Limpio",
        data=excel_buffer.getvalue(),
        file_name="excel_limpio.xlsx"
    )
